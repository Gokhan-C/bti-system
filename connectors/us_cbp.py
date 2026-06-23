"""
ABD CBP Connector — v3

Veri çekme stratejisi: YENİ KARAR TESPİTİ
  - CBP API'den en son yayınlanan kararlar çekilir.
  - state/us_cbp_seen.json ile karşılaştırılır → sadece yeni ruling numaraları işlenir.
  - Kararlar kategorize edilir: 'classification', 'origin', 'other'
  - Rapora SADECE 'classification' kararları girer.
  - Claude ile token-optimize 3 madde özet üretilir.
  - Her kararın başında tam metne erişim linki verilir.
  - Raporun başında günlük istatistik tablosu yer alır.
  - _report_data birleşik rapor için saklanır.
"""

import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import requests

from core.base_connector import BaseConnector
from core.report_builder import (
    make_doc, add_info_table, add_ruling_link,
    add_summary_section, add_no_results_notice, set_cell_bg,
)
from core.translator import translate_google

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


CBP_API_RECENT = "https://rulings.cbp.gov/api/stat/recentRulings"
CBP_API_DETAIL = "https://rulings.cbp.gov/api/ruling/{number}"

HEADERS = {
    "Accept": "application/json, text/plain, */*",
    "Referer": "https://rulings.cbp.gov/home",
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
}

ORIGIN_KEYWORDS = [
    "COUNTRY OF ORIGIN", "COUNTRY-OF-ORIGIN",
    "NAFTA", "USMCA", "CAFTA", "FREE TRADE AGREEMENT",
    "RULES OF ORIGIN", "SUBSTANTIAL TRANSFORMATION",
    "MARKING", "CATEGORY: ORIGIN", "CATEGORY:  ORIGIN",
]

CLASSIFICATION_KEYWORDS = [
    "TARIFF CLASSIFICATION", "HTS ", "HTSUS ", "HARMONIZED TARIFF",
    "CATEGORY: CLASSIFICATION", "CATEGORY:  CLASSIFICATION",
]


def _extract_product_desc(text: str, max_chars: int = 1800) -> str:
    """
    CBP karar metninden ürün tanımı bölümünü çıkarır.
    'In your letter dated...' ile 'The applicable subheading...' arasındaki
    paragraf ürün adı, marka/model ve teknik özellikleri içerir.
    """
    if not text:
        return ""
    upper = text.upper()
    start_idx = -1
    for marker in (
        "IN YOUR LETTER",
        "THE ITEMS CONCERNED",
        "THE ITEM UNDER CONSIDERATION",
        "THE MERCHANDISE UNDER CONSIDERATION",
        "THE SUBJECT MERCHANDISE",
    ):
        idx = upper.find(marker)
        if idx != -1:
            start_idx = idx
            break

    end_idx = upper.find("THE APPLICABLE SUBHEADING")
    if end_idx == -1:
        end_idx = upper.find("THE APPLICABLE HTS")
    if end_idx == -1:
        end_idx = upper.find("HARMONIZED TARIFF SCHEDULE")

    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        return text[start_idx:end_idx].strip()[:max_chars]
    if start_idx != -1:
        return text[start_idx: start_idx + max_chars].strip()
    return text[:max_chars].strip()


def _extract_classification_rationale(text: str, max_chars: int = 1500) -> str:
    """
    CBP karar metninden sınıflandırma gerekçesini çıkarır.
    'The applicable subheading for [ürün] will be XXXX.XX, HTSUS, which provides for...'
    paragrafı; hangi tarife pozisyonunun neden uygulandığını açıklar.
    """
    if not text:
        return ""
    upper = text.upper()
    idx = upper.find("THE APPLICABLE SUBHEADING FOR")
    if idx == -1:
        idx = upper.find("THE APPLICABLE SUBHEADING")
    if idx == -1:
        idx = upper.find("THE CORRECT SUBHEADING")

    if idx != -1:
        snippet = text[idx: idx + max_chars]
        for cutoff in (
            "The duties cited above",
            "This ruling does not address",
            "The holding set forth",
            "For further information",
        ):
            ci = snippet.find(cutoff)
            if ci != -1:
                snippet = snippet[:ci]
        return snippet.strip()
    return ""


def _summarize_ruling(
    text: str,
    subject: str,
    tariffs: str,
    ruling_number: str,
    logger=None,
) -> dict[str, str]:
    """
    CBP karar metninden ürün tanımı ve sınıflandırma gerekçesini çıkarıp
    Google Translate ile Türkçeye çevirir.
    """
    prod_text = _extract_product_desc(text)
    rat_text  = _extract_classification_rationale(text)

    esya_tanimi = translate_google(prod_text or subject, logger=logger)
    gtip_karar  = tariffs
    teknik_gerekce = (
        translate_google(rat_text, logger=logger) if rat_text
        else "(Gerekçe metni bulunamadı)"
    )

    return {
        "esya_tanimi":    esya_tanimi or subject,
        "gtip_karar":     gtip_karar,
        "teknik_gerekce": teknik_gerekce,
    }


def _categorize_ruling(detail: dict) -> str:
    subject   = (detail.get("subject") or "").upper()
    text_head = (detail.get("text") or "")[:600].upper()
    has_tariffs = bool(detail.get("tariffs"))

    is_origin = any(kw in subject or kw in text_head for kw in ORIGIN_KEYWORDS)
    is_class  = has_tariffs or any(kw in subject or kw in text_head for kw in CLASSIFICATION_KEYWORDS)

    if is_origin and not is_class:
        return "origin"
    if is_class:
        return "classification"
    return "other"


def _fetch_recent_rulings(recent_days: int) -> list[dict]:
    try:
        r = requests.get(
            CBP_API_RECENT,
            params={"format": "json", "collection": ""},
            headers=HEADERS, timeout=20,
        )
        r.raise_for_status()
        cutoff = (datetime.now() - timedelta(days=recent_days)).strftime("%Y-%m-%d")
        result = []
        for ruling in r.json():
            try:
                dt = datetime.strptime(ruling.get("dateModified", ""), "%m/%d/%Y")
                if dt.strftime("%Y-%m-%d") >= cutoff:
                    result.append(ruling)
            except Exception:
                pass
        return result
    except Exception:
        return []


def _fetch_detail(ruling_number: str) -> dict:
    try:
        r = requests.get(CBP_API_DETAIL.format(number=ruling_number), headers=HEADERS, timeout=20)
        r.raise_for_status()
        return r.json()
    except Exception:
        return {}


def _add_stats_table(doc, stats: dict, date_str: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Günlük Karar İstatistikleri")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    add_info_table(doc, [
        ("Tarih",                         date_str),
        ("Toplam Çekilen Karar",          str(stats["total"])),
        ("✓ Tarife Sınıflandırması",      str(stats["classification"])),
        ("✗ Menşei Kararı (raporda yok)", str(stats["origin"])),
        ("– Diğer",                       str(stats["other"])),
        ("Rapora Giren",                  str(stats["in_report"])),
    ])
    doc.add_paragraph()


class UsCbpConnector(BaseConnector):

    @property
    def connector_id(self) -> str:
        return "us_cbp"

    @property
    def display_name(self) -> str:
        return self.config.get("display_name", "ABD CBP")

    def extract_id(self, record: dict[str, Any]) -> str:
        return str(record.get("rulingNumber") or record.get("id") or "")

    def fetch(self, target_date: datetime) -> list[dict[str, Any]]:
        recent_days = self.config.get("recent_days", 14)
        rulings = _fetch_recent_rulings(recent_days)
        for ruling in rulings:
            ruling["id"] = ruling.get("rulingNumber", "")
        return rulings

    def build_report(self, records: list[dict[str, Any]], target_date: datetime) -> list[Path]:
        date_str = target_date.strftime("%Y-%m-%d")
        date_dir = self.output_dir / date_str
        date_dir.mkdir(parents=True, exist_ok=True)

        doc = make_doc("ABD CBP Tarife Sınıflandırma Kararları", f"CBP  |  {date_str}")

        stats = {"total": 0, "classification": 0, "origin": 0, "other": 0, "in_report": 0}
        processed_records = []

        for ruling in records:
            number = ruling.get("rulingNumber", "UNKNOWN")
            detail = _fetch_detail(number)
            if not detail:
                continue

            stats["total"] += 1
            category = _categorize_ruling(detail)
            stats[category] += 1

            if category != "classification":
                time.sleep(0.3)
                continue

            date_raw = ruling.get("dateModified", date_str)
            try:
                date_fmt = datetime.strptime(date_raw, "%m/%d/%Y").strftime("%Y-%m-%d")
            except Exception:
                date_fmt = date_str

            tariffs   = ", ".join(detail.get("tariffs") or []) or "-"
            subject   = detail.get("subject") or "-"
            text      = detail.get("text") or ""
            collection = (detail.get("collection") or ruling.get("collection") or "").upper()
            source_url = f"https://rulings.cbp.gov/ruling/{number}"

            summary = _summarize_ruling(
                text=text,
                subject=subject,
                tariffs=tariffs,
                ruling_number=number,
            )

            processed_records.append({
                "number":     number,
                "date_fmt":   date_fmt,
                "collection": collection,
                "tariffs":    tariffs,
                "summary":    summary,
                "source_url": source_url,
            })
            time.sleep(0.3)

        stats["in_report"] = len(processed_records)
        _add_stats_table(doc, stats, date_str)

        for rec in processed_records:
            add_ruling_link(doc, rec["source_url"])

            hp = doc.add_paragraph()
            hr = hp.add_run(f"ABD CBP Kararı: {rec['number']}  |  {rec['date_fmt']}")
            hr.bold = True
            hr.font.size = Pt(14)
            hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            add_info_table(doc, [
                ("Karar Numarası", rec["number"]),
                ("Koleksiyon",     rec["collection"]),
                ("Karar Tarihi",   rec["date_fmt"]),
                ("HTS / GTİP",     rec["tariffs"]),
            ])
            doc.add_paragraph()
            add_summary_section(doc, rec["summary"])

        if not processed_records:
            add_no_results_notice(doc)

        # Birleşik rapor için veriyi sakla
        self._report_data = {
            "stats":   stats,
            "records": processed_records,
            "date_str": date_str,
        }
        # JSON'a kalıcı kaydet — records_new=0 olduğunda orchestrator buradan okur
        import json as _json
        (date_dir / "_report_data.json").write_text(
            _json.dumps(self._report_data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        docx_path = date_dir / f"ABD_CBP_Kararlar_{date_str}.docx"
        doc.save(str(docx_path))
        return [docx_path]
