"""
Kanada CBSA Connector — v3

Veri çekme stratejisi: YENİ KARAR TESPİTİ
  - CBSA CARM API'den ruling listesi çekilir.
  - state/ca_cbsa_seen.json ile karşılaştırılır → sadece yeni rulingsId'ler işlenir.
  - Sadece HS tarife sınıflandırması kararları rapora girer.
  - Claude ile token-optimize 3 madde özet üretilir.
  - Her kararın başında tam metne erişim linki verilir.
  - _report_data birleşik rapor için saklanır.
"""

import time
from datetime import datetime
from pathlib import Path
from typing import Any

import requests

from core.base_connector import BaseConnector
from core.report_builder import (
    make_doc, add_info_table, add_ruling_link,
    add_summary_section, add_no_results_notice,
)
from core.translator import summarize_ruling_claude


CBSA_LIST_URL    = "https://ccp-pcc.cbsa-asfc.cloud-nuage.canada.ca/carmwebservices/v2/carm/publicrulings/list"
CBSA_DETAIL_URL  = "https://ccp-pcc.cbsa-asfc.cloud-nuage.canada.ca/carmwebservices/v2/carm/publicrulings/rulingsdetails/{id}"

HEADERS = {
    "Accept": "application/json, text/plain, */*",
    "Referer": "https://ccp-pcc.cbsa-asfc.cloud-nuage.canada.ca/en/national-rulings",
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
}

HS_KEYWORDS = [
    "TARIFF CLASSIFICATION", "HTS", "HARMONIZED", "CLASSIFICATION",
]

from docx.shared import Pt, RGBColor


def _is_hs_classification(detail: dict) -> bool:
    if detail.get("tariffs"):
        return True
    text    = (detail.get("text") or detail.get("analysisAndJustification") or "")[:600].upper()
    subject = (detail.get("subject") or detail.get("decision") or "")[:300].upper()
    return any(kw in text or kw in subject for kw in HS_KEYWORDS)


def _fetch_list(api_url: str) -> list[dict]:
    try:
        r = requests.get(api_url, params={"lang": "en", "curr": "CAD"}, headers=HEADERS, timeout=20)
        r.raise_for_status()
        data = r.json()
        return data if isinstance(data, list) else data.get("items", data.get("rulings", []))
    except Exception:
        return []


def _fetch_detail(ruling_id: str, detail_url_tpl: str) -> dict:
    try:
        r = requests.get(detail_url_tpl.format(id=ruling_id) + "?lang=en&curr=CAD", headers=HEADERS, timeout=20)
        r.raise_for_status()
        return r.json()
    except Exception:
        return {}


class CaCbsaConnector(BaseConnector):

    @property
    def connector_id(self) -> str:
        return "ca_cbsa"

    @property
    def display_name(self) -> str:
        return self.config.get("display_name", "Kanada CBSA")

    def extract_id(self, record: dict[str, Any]) -> str:
        return str(
            record.get("rulingsId") or record.get("rulingId")
            or record.get("id") or record.get("caseNumber") or ""
        )

    def fetch(self, target_date: datetime) -> list[dict[str, Any]]:
        api_url = self.config.get("api_url", CBSA_LIST_URL)
        return _fetch_list(api_url)

    def build_report(self, records: list[dict[str, Any]], target_date: datetime) -> list[Path]:
        date_str        = target_date.strftime("%Y-%m-%d")
        date_dir        = self.output_dir / date_str
        date_dir.mkdir(parents=True, exist_ok=True)
        detail_url_tpl  = self.config.get("api_detail_url", CBSA_DETAIL_URL)

        doc = make_doc("Kanada CBSA Tarife Sınıflandırma Kararları", f"CBSA  |  {date_str}")
        processed_records = []

        for ruling in records:
            ruling_id   = self.extract_id(ruling)
            date_fmt    = (ruling.get("date") or ruling.get("dateOfDecision") or date_str)[:10]
            product     = (ruling.get("product") or ruling.get("productName") or ruling.get("description") or "-")
            hts         = ruling.get("classificationNumber") or ruling.get("tariffCode") or "-"
            ruling_type = ruling.get("typeOfRuling") or ruling.get("rulingType") or "-"

            detail   = _fetch_detail(ruling_id, detail_url_tpl)
            if not _is_hs_classification(detail):
                time.sleep(0.3)
                continue

            analysis  = detail.get("analysisAndJustification") or detail.get("analysis") or ""
            decision  = detail.get("decision") or "-"
            applicant = detail.get("applicantName") or detail.get("applicant") or "-"
            origin    = detail.get("countryOfOrigin") or detail.get("origin") or "-"
            prod_desc = detail.get("productDescription") or detail.get("description") or product
            source_url = (
                f"https://ccp-pcc.cbsa-asfc.cloud-nuage.canada.ca"
                f"/en/national-rulings/national-rulings-details/{ruling_id}"
            )

            summary = summarize_ruling_claude(
                product_desc=prod_desc,
                analysis_text=analysis,
                decision=decision,
                gtip_codes=hts,
                ruling_number=ruling_id,
            )

            processed_records.append({
                "ruling_id":   ruling_id,
                "ruling_type": ruling_type,
                "date_fmt":    date_fmt,
                "hts":         hts,
                "applicant":   applicant,
                "origin":      origin,
                "summary":     summary,
                "source_url":  source_url,
            })
            time.sleep(0.3)

        for rec in processed_records:
            add_ruling_link(doc, rec["source_url"])

            hp = doc.add_paragraph()
            hr = hp.add_run(f"Kanada CBSA Kararı: {rec['ruling_id']}  |  {rec['date_fmt']}")
            hr.bold = True
            hr.font.size = Pt(14)
            hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            add_info_table(doc, [
                ("Karar Numarası", rec["ruling_id"]),
                ("Karar Türü",     rec["ruling_type"]),
                ("Karar Tarihi",   rec["date_fmt"]),
                ("GTİP Kodu",      rec["hts"]),
                ("Başvurucu",      rec["applicant"]),
                ("Menşe Ülke",     rec["origin"]),
            ])
            doc.add_paragraph()
            add_summary_section(doc, rec["summary"])

        if not processed_records:
            add_no_results_notice(doc)

        # Birleşik rapor için veriyi sakla
        self._report_data = {
            "records":  processed_records,
            "date_str": date_str,
        }
        # JSON'a kalıcı kaydet — records_new=0 olduğunda orchestrator buradan okur
        import json as _json
        (date_dir / "_report_data.json").write_text(
            _json.dumps(self._report_data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        docx_path = date_dir / f"Kanada_CBSA_Kararlar_{date_str}.docx"
        doc.save(str(docx_path))
        return [docx_path]
