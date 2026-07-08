"""
İngiltere HMRC Connector — Advance Tariff Rulings (ATaR)

Veri kaynağı: https://www.tax.service.gov.uk/search-for-advance-tariff-rulings
  - JSON API yok → HTML scraping (govuk-summary-list yapısı, bs4 ile parse).
  - Boş arama tüm kararları yeniden-eskiye listeler (sayfa başı 25 kayıt).

Veri çekme stratejisi: YENİ KARAR TESPİTİ (ca_cbsa modeli)
  - İlk `search_pages` sayfadan ruling ID'leri toplanır.
  - state/uk_hmrc_seen.json ile karşılaştırılır → sadece yeni ID'ler işlenir.
  - Start date `recent_days` dışındaysa rapora girmez (ID yine seen'e yazılır)
    → ilk çalıştırmada eski kararlarla taşma olmaz.
  - Claude ile token-optimize 3 madde özet üretilir.
  - Her kararın başında tam metne erişim linki verilir (detay sayfası GET ile açılır).
  - _report_data.json birleşik rapor ve site için saklanır.
"""

import re
import time
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import requests
from bs4 import BeautifulSoup
from docx.shared import Pt, RGBColor

from core.base_connector import BaseConnector
from core.report_builder import (
    make_doc, add_info_table, add_ruling_link,
    add_summary_section, add_no_results_notice,
)
from core.translator import summarize_ruling_claude


UK_BASE_URL = "https://www.tax.service.gov.uk/search-for-advance-tariff-rulings"

HEADERS = {
    "Accept": "text/html,application/xhtml+xml",
    "Accept-Language": "en-GB,en;q=0.9",
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
}


def _fetch_search_page(base_url: str, page: int) -> list[str]:
    """Arama sayfasındaki ruling ID'lerini (görünüm sırasıyla) döndürür."""
    try:
        # Not: boş "query=" parametresi verilirse servis her zaman 1. sayfayı döndürür;
        # sayfalama linkleri gibi yalnızca "page" gönderilmeli.
        r = requests.get(
            f"{base_url}/search", params={"page": page},
            headers=HEADERS, timeout=30,
        )
        r.raise_for_status()
    except Exception:
        return []
    ids = re.findall(r'href="[^"]*/ruling/(\d+)"', r.text)
    # Sıra korunarak tekilleştir
    seen: set[str] = set()
    return [i for i in ids if not (i in seen or seen.add(i))]


def _parse_uk_date(s: str) -> str:
    """'14 Jul 2023' → '2023-07-14' (çözülemezse ham metin)."""
    try:
        return datetime.strptime(s.strip(), "%d %b %Y").strftime("%Y-%m-%d")
    except Exception:
        return s.strip()


def _fetch_detail(base_url: str, ruling_id: str) -> dict:
    """Detay sayfasındaki govuk-summary-list alanlarını sözlüğe çevirir."""
    try:
        r = requests.get(f"{base_url}/ruling/{ruling_id}", headers=HEADERS, timeout=30)
        r.raise_for_status()
    except Exception:
        return {}

    soup = BeautifulSoup(r.text, "html.parser")
    fields: dict[str, str] = {}
    for row in soup.select(".govuk-summary-list__row"):
        key = row.select_one(".govuk-summary-list__key")
        val = row.select_one(".govuk-summary-list__value")
        if key and val:
            fields[key.get_text(strip=True).lower()] = val.get_text(" ", strip=True)
    return fields


class UkHmrcConnector(BaseConnector):

    @property
    def connector_id(self) -> str:
        return "uk_hmrc"

    @property
    def display_name(self) -> str:
        return self.config.get("display_name", "İngiltere HMRC")

    def extract_id(self, record: dict[str, Any]) -> str:
        return str(record.get("id", ""))

    def fetch(self, target_date: datetime) -> list[dict[str, Any]]:
        base_url = self.config.get("base_url", UK_BASE_URL)
        pages    = int(self.config.get("search_pages", 3))
        records: list[dict[str, Any]] = []
        collected: set[str] = set()
        for page in range(1, pages + 1):
            for rid in _fetch_search_page(base_url, page):
                if rid not in collected:
                    collected.add(rid)
                    records.append({"id": rid})
            time.sleep(0.5)
        return records

    def build_report(self, records: list[dict[str, Any]], target_date: datetime) -> list[Path]:
        base_url    = self.config.get("base_url", UK_BASE_URL)
        recent_days = int(self.config.get("recent_days", 14))
        cutoff      = (target_date - timedelta(days=recent_days)).strftime("%Y-%m-%d")

        date_str = target_date.strftime("%Y-%m-%d")
        date_dir = self.output_dir / date_str
        date_dir.mkdir(parents=True, exist_ok=True)

        doc = make_doc("İngiltere HMRC Tarife Sınıflandırma Kararları (ATaR)",
                       f"HMRC  |  {date_str}")
        processed_records = []

        for rec in records:
            ruling_id = self.extract_id(rec)
            detail    = _fetch_detail(base_url, ruling_id)
            time.sleep(0.5)
            if not detail:
                continue

            start_iso = _parse_uk_date(detail.get("start date", ""))
            # Eski karar (ilk çalıştırma taşması) → rapora alma, ID seen'de kalır
            if start_iso and start_iso < cutoff:
                continue

            expiry_iso = _parse_uk_date(detail.get("expiry date", ""))
            commodity  = re.sub(r"\s*\(opens in new tab\)\s*", "", detail.get("commodity code", "-")).strip() or "-"
            desc       = detail.get("description", "")
            just       = detail.get("justification", "")
            keywords   = detail.get("keywords", "")
            source_url = f"{base_url}/ruling/{ruling_id}"

            summary = summarize_ruling_claude(
                product_desc=desc,
                analysis_text=just,
                decision="",
                gtip_codes=commodity,
                ruling_number=ruling_id,
            )

            processed_records.append({
                "ruling_id":  ruling_id,
                "date_fmt":   start_iso or date_str,
                "expiry":     expiry_iso,
                "hts":        commodity,
                "keywords":   keywords,
                "summary":    summary,
                "source_url": source_url,
            })

        # Aynı gün önceki koşunun kayıtlarıyla birleştir (retry/ikinci koşu veriyi ezmesin)
        import json as _json
        json_path = date_dir / "_report_data.json"
        if json_path.exists():
            try:
                old = _json.loads(json_path.read_text(encoding="utf-8")).get("records", [])
                have = {r["ruling_id"] for r in processed_records}
                merged_old = [r for r in old if r.get("ruling_id") not in have]
                processed_records = sorted(
                    processed_records + merged_old,
                    key=lambda r: r.get("date_fmt", ""), reverse=True,
                )
            except Exception:
                pass

        for rec in processed_records:
            add_ruling_link(doc, rec["source_url"])

            hp = doc.add_paragraph()
            hr = hp.add_run(f"İngiltere HMRC Kararı: {rec['ruling_id']}  |  {rec['date_fmt']}")
            hr.bold = True
            hr.font.size = Pt(14)
            hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            add_info_table(doc, [
                ("Karar Numarası",      rec["ruling_id"]),
                ("Karar Tarihi",        rec["date_fmt"]),
                ("Geçerlilik Bitişi",   rec["expiry"] or "-"),
                ("GTİP Kodu",           rec["hts"]),
                ("Anahtar Kelimeler",   rec["keywords"] or "-"),
            ])
            doc.add_paragraph()
            add_summary_section(doc, rec["summary"])

        if not processed_records:
            add_no_results_notice(doc)

        # Birleşik rapor + site için veriyi sakla
        self._report_data = {
            "records":  processed_records,
            "date_str": date_str,
        }
        json_path.write_text(
            _json.dumps(self._report_data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        docx_path = date_dir / f"Ingiltere_HMRC_Kararlar_{date_str}.docx"
        doc.save(str(docx_path))
        return [docx_path]
