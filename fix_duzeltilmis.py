#!/usr/bin/env python3
"""
Düzeltme scripti: EU EBTI ve ABD CBP raporlarını yeniden oluşturur.
Çeviri: Google Translate (deep-translator, source=auto).
Dosya isimlerine '_düzeltilmiş' eki eklenir.
"""

import csv
import json
import subprocess
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

import requests
from deep_translator import GoogleTranslator
from docx.shared import Pt, RGBColor

from core.report_builder import (
    make_doc, add_info_table, add_ruling_link,
    add_summary_section, add_no_results_notice,
)

# ── Sabitler ──────────────────────────────────────────────────────────────────

BASE        = Path.home() / "BTI_Reports"
BTI_DATE    = "2026-05-20"
EU_DIR      = BASE / "EU_EBTI" / BTI_DATE
US_DIR      = BASE / "US_CBP"  / BTI_DATE
CA_DIR      = BASE / "CA_CBSA" / BTI_DATE

DOCX_SCRIPT         = Path.home() / "bti_system/assets/build_docx.js"
UNIFIED_DOCX_SCRIPT = Path.home() / "bti_system/assets/build_unified_docx.js"

CBP_HEADERS = {
    "Accept": "application/json, text/plain, */*",
    "Referer": "https://rulings.cbp.gov/home",
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
}

COUNTRY_COLORS = {
    "eu_ebti": {"color": "003399", "label_color": "2E75B6"},
    "us_cbp":  {"color": "B22234", "label_color": "C0392B"},
    "ca_cbsa": {"color": "FF0000", "label_color": "C0392B"},
}

CONNECTOR_DISPLAY = {
    "eu_ebti": "Avrupa Birliği EBTI",
    "us_cbp":  "ABD CBP",
    "ca_cbsa": "Kanada CBSA",
}

CHUNK_SIZE = 4500  # Google Translate karakter limiti


# ── Google Translate yardımcısı ───────────────────────────────────────────────

def google_tr(text: str, source: str = "auto") -> str:
    """Metni Türkçeye çevirir. Kaynak dil otomatik tespitli."""
    if not text or not text.strip():
        return text
    try:
        translator = GoogleTranslator(source=source, target="tr")
        chunks = [text[i: i + CHUNK_SIZE] for i in range(0, len(text), CHUNK_SIZE)]
        parts = []
        for chunk in chunks:
            parts.append(translator.translate(chunk))
            time.sleep(0.25)
        return " ".join(p for p in parts if p)
    except Exception as e:
        print(f"    Google Translate hatası: {e}")
        return text


# ── EU EBTI ───────────────────────────────────────────────────────────────────

def fix_eu_ebti() -> dict:
    print("\n=== EU EBTI düzeltmesi başlıyor ===")

    csv_path  = EU_DIR / f"BTI_{BTI_DATE.replace('-', '')}.csv"
    json_path = EU_DIR / "_report_data.json"

    with open(csv_path, encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))
    print(f"  {len(rows)} kayıt CSV'den okundu")

    # ref → {desc, just, lang} eşlemesi
    ref_map: dict[str, dict] = {}
    for row in rows:
        ref = (
            row.get("BTI_REFERENCE") or row.get("BTI Reference")
            or row.get("Reference") or ""
        ).strip()
        if ref:
            ref_map[ref] = {
                "desc": (row.get("DESCRIPTION_OF_GOODS") or "").strip(),
                "just": (row.get("CLASSIFICATION_JUSTIFICATION") or "").strip(),
                "lang": (row.get("LANGUAGE") or "auto").strip().lower(),
            }

    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    total = len(data["records"])
    print(f"  {total} kayıt Google Translate ile Türkçeye çevriliyor...")

    for i, rec in enumerate(data["records"], 1):
        raw = ref_map.get(rec["ref"], {})
        desc_orig = raw.get("desc", rec.get("desc_tr", ""))
        just_orig = raw.get("just", rec.get("just_tr", ""))
        lang      = raw.get("lang", "auto")

        if i % 10 == 1:
            print(f"  [{i}/{total}] işleniyor...")

        rec["desc_tr"] = google_tr(desc_orig, source=lang) if desc_orig else desc_orig
        rec["just_tr"] = google_tr(just_orig, source=lang) if just_orig else just_orig

    print(f"  Çeviri tamamlandı. _report_data.json güncelleniyor...")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    # Node.js ile yeni docx üret
    tmp_json = EU_DIR / "_eu_duzeltilmis_tmp.json"
    with open(tmp_json, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    date_tr      = data.get("report_date_tr", "20Mayıs2026")
    date_compact = date_tr.replace(" ", "")
    docx_path    = EU_DIR / f"AB_EBTI_Rapor_{date_compact}_düzeltilmiş.docx"

    result = subprocess.run(
        ["node", str(DOCX_SCRIPT), str(tmp_json), str(docx_path)],
        capture_output=True, text=True,
    )
    tmp_json.unlink(missing_ok=True)

    if result.returncode != 0:
        print(f"  HATA (Node.js): {result.stderr[:400]}")
    else:
        print(f"  ✓ Kaydedildi: {docx_path.name}")

    return data


# ── ABD CBP ───────────────────────────────────────────────────────────────────

def _fetch_cbp_detail(ruling_number: str) -> dict:
    try:
        r = requests.get(
            f"https://rulings.cbp.gov/api/ruling/{ruling_number}",
            headers=CBP_HEADERS, timeout=25,
        )
        r.raise_for_status()
        return r.json()
    except Exception as e:
        print(f"    API hatası ({ruling_number}): {e}")
        return {}


def _extract_product_desc(text: str, max_chars: int = 1800) -> str:
    """
    Ürün tanımı bölümünü çıkarır.
    CBP kararlarında ürün tanımı genellikle 'In your letter dated' ile
    'The applicable subheading' arasında yer alır.
    """
    if not text:
        return ""
    upper = text.upper()

    # Başlangıç: 'IN YOUR LETTER' veya 'THE ITEM' veya 'THE MERCHANDISE'
    start_idx = -1
    for marker in ("IN YOUR LETTER", "THE ITEMS CONCERNED", "THE ITEM UNDER CONSIDERATION",
                   "THE MERCHANDISE UNDER CONSIDERATION", "THE SUBJECT MERCHANDISE"):
        idx = upper.find(marker)
        if idx != -1:
            start_idx = idx
            break

    # Bitiş: 'THE APPLICABLE SUBHEADING' (sınıflandırma kararının başladığı yer)
    end_idx = upper.find("THE APPLICABLE SUBHEADING")
    if end_idx == -1:
        end_idx = upper.find("THE APPLICABLE HTS")
    if end_idx == -1:
        end_idx = upper.find("HARMONIZED TARIFF SCHEDULE")

    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        return text[start_idx:end_idx].strip()[:max_chars]
    if start_idx != -1:
        return text[start_idx: start_idx + max_chars].strip()
    return text[:max_chars].strip()


def _extract_classification_rationale(text: str, max_chars: int = 1500) -> str:
    """
    Sınıflandırma gerekçesini çıkarır.
    'The applicable subheading for [ürün adı] will be XXXX.XX.XXXX, HTSUS,
    which provides for ...' paragrafını alır.
    """
    if not text:
        return ""
    upper = text.upper()

    # 'THE APPLICABLE SUBHEADING FOR...' paragrafını bul
    idx = upper.find("THE APPLICABLE SUBHEADING FOR")
    if idx == -1:
        idx = upper.find("THE APPLICABLE SUBHEADING")
    if idx == -1:
        idx = upper.find("THE CORRECT SUBHEADING")

    if idx != -1:
        # Bu paragrafı al; bir sonraki boş satır veya max_chars ile kes
        snippet = text[idx: idx + max_chars]
        # Boilerplate başlangıcını kes (duty rates, disclaimers)
        for cutoff in ("The duties cited above", "This ruling does not address",
                       "The holding set forth", "For further information"):
            ci = snippet.find(cutoff)
            if ci != -1:
                snippet = snippet[:ci]
        return snippet.strip()

    return ""


def _add_cbp_stats_table(doc, stats: dict, date_str: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Günlük Karar İstatistikleri")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    add_info_table(doc, [
        ("Tarih",                         date_str),
        ("Toplam Çekilen Karar",          str(stats.get("total", 0))),
        ("✓ Tarife Sınıflandırması",      str(stats.get("classification", 0))),
        ("✗ Menşei Kararı (raporda yok)", str(stats.get("origin", 0))),
        ("– Diğer",                       str(stats.get("other", 0))),
        ("Rapora Giren",                  str(stats.get("in_report", 0))),
    ])
    doc.add_paragraph()


def fix_us_cbp() -> dict:
    print("\n=== ABD CBP düzeltmesi başlıyor ===")

    json_path = US_DIR / "_report_data.json"
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    records  = data["records"]
    stats    = data["stats"]
    date_str = data["date_str"]

    print(f"  {len(records)} sınıflandırma kararı işlenecek")

    fixed_records = []
    for i, rec in enumerate(records, 1):
        number = rec["number"]
        print(f"  [{i}/{len(records)}] {number} işleniyor...")

        detail  = _fetch_cbp_detail(number)
        subject = detail.get("subject") or ""
        text    = detail.get("text") or ""
        tariffs = rec["tariffs"]

        # Eşya tanımı: ürün tanımı paragrafını çıkar + Türkçeye çevir
        prod_text   = _extract_product_desc(text)
        esya_tanimi = google_tr(prod_text, source="en") if prod_text else google_tr(subject, source="en")

        # GTİP kararı
        gtip_karar = tariffs

        # Teknik gerekçe: sınıflandırma gerekçesi paragrafını çıkar + Türkçeye çevir
        rat_text       = _extract_classification_rationale(text)
        teknik_gerekce = google_tr(rat_text, source="en") if rat_text else "(Gerekçe metni bulunamadı)"

        summary = {
            "esya_tanimi":   esya_tanimi,
            "gtip_karar":    gtip_karar,
            "teknik_gerekce": teknik_gerekce,
        }
        fixed_records.append({**rec, "summary": summary})
        time.sleep(0.3)

    data["records"] = fixed_records

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print("  _report_data.json güncellendi")

    # python-docx ile yeni rapor üret
    doc = make_doc("ABD CBP Tarife Sınıflandırma Kararları", f"CBP  |  {date_str}")
    _add_cbp_stats_table(doc, stats, date_str)

    for rec in fixed_records:
        add_ruling_link(doc, rec["source_url"])

        hp = doc.add_paragraph()
        hr = hp.add_run(f"ABD CBP Kararı: {rec['number']}  |  {rec['date_fmt']}")
        hr.bold = True
        hr.font.size = Pt(14)
        hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        add_info_table(doc, [
            ("Karar Numarası", rec["number"]),
            ("Koleksiyon",     rec["collection"]),
            ("Karar Tarihi",   rec["date_fmt"]),
            ("HTS / GTİP",     rec["tariffs"]),
        ])
        doc.add_paragraph()
        add_summary_section(doc, rec["summary"])

    if not fixed_records:
        add_no_results_notice(doc)

    docx_path = US_DIR / f"ABD_CBP_Kararlar_{date_str}_düzeltilmiş.docx"
    doc.save(str(docx_path))
    print(f"  ✓ Kaydedildi: {docx_path.name}")

    return data


# ── Birleşik Rapor ────────────────────────────────────────────────────────────

def fix_unified(eu_data: dict, us_data: dict) -> None:
    print("\n=== Birleşik rapor düzeltmesi başlıyor ===")

    ca_data = None
    ca_json = CA_DIR / "_report_data.json"
    if ca_json.exists():
        with open(ca_json, encoding="utf-8") as f:
            ca_data = json.load(f)

    country_cards = []
    for cid, display in CONNECTOR_DISPLAY.items():
        colors = COUNTRY_COLORS.get(cid, {"color": "444444", "label_color": "666666"})
        if cid == "eu_ebti":
            count = len(eu_data.get("records", []))
        elif cid == "us_cbp":
            count = len(us_data.get("records", []))
        elif cid == "ca_cbsa" and ca_data:
            count = len(ca_data.get("records", []))
        else:
            count = 0
        country_cards.append({
            "name":        display,
            "count":       count,
            "color":       colors["color"],
            "label_color": colors["label_color"],
        })

    unified_data: dict = {
        "date":      BTI_DATE,
        "countries": country_cards,
        "eu_ebti":   eu_data,
        "us_cbp":    {**us_data, "date_str": BTI_DATE},
    }
    if ca_data:
        unified_data["ca_cbsa"] = {**ca_data, "date_str": BTI_DATE}

    tmp_json  = BASE / "_unified_duzeltilmis_tmp.json"
    docx_path = BASE / f"BTI_Unified_{BTI_DATE}_düzeltilmiş.docx"

    tmp_json.write_text(
        json.dumps(unified_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    result = subprocess.run(
        ["node", str(UNIFIED_DOCX_SCRIPT), str(tmp_json), str(docx_path)],
        capture_output=True, text=True,
    )
    tmp_json.unlink(missing_ok=True)

    if result.returncode != 0:
        print(f"  HATA (Node.js): {result.stderr[:400]}")
    else:
        print(f"  ✓ Kaydedildi: {docx_path.name}")


# ── Ana ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    eu_data = fix_eu_ebti()
    us_data = fix_us_cbp()
    fix_unified(eu_data, us_data)

    print("\n✓ Tüm düzeltmeler tamamlandı.")
    print(f"  • EU_EBTI/{BTI_DATE}/AB_EBTI_Rapor_20Mayıs2026_düzeltilmiş.docx")
    print(f"  • US_CBP/{BTI_DATE}/ABD_CBP_Kararlar_{BTI_DATE}_düzeltilmiş.docx")
    print(f"  • BTI_Unified_{BTI_DATE}_düzeltilmiş.docx")
