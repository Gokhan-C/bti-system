"""
Ortak python-docx yardımcıları — tüm connector'lar tarafından kullanılır.
AB EBTI connector'ı bireysel raporunu Node.js ile üretir, birleşik rapor için burası kullanılır.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Temel yapılar ─────────────────────────────────────────────────────────────

def make_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.5))

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(17)
    tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    return doc


def set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        lc.width = Cm(5)
        vc.width = Cm(11)
        set_cell_bg(lc, "EBF3FB")
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        vr = vc.paragraphs[0].add_run(str(value) if value else "-")
        vr.font.size = Pt(10)


def add_ruling_to_doc(
    doc: Document,
    heading: str,
    info_rows: list[tuple[str, str]],
    sections: list[tuple[str, str | list]],
) -> None:
    hp = doc.add_paragraph()
    hr = hp.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_info_table(doc, info_rows)
    doc.add_paragraph()

    for sec_title, content in sections:
        sp = doc.add_paragraph()
        sr = sp.add_run(sec_title)
        sr.bold = True
        sr.font.size = Pt(11)
        sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        if isinstance(content, str) and content:
            p = doc.add_paragraph()
            p.add_run(content).font.size = Pt(10)
        elif isinstance(content, list):
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item).font.size = Pt(10)
        doc.add_paragraph()

    doc.add_paragraph("─" * 80)
    doc.add_paragraph()


def add_no_results_notice(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bugün yeni HS sınıflandırma kararı bulunamadı.")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── Karar kartı bileşenleri (CBP + CBSA ortak) ────────────────────────────────

def add_ruling_link(doc: Document, url: str, label: str = "Tam metne erişmek için tıklayın") -> None:
    """Kararın başına tıklanabilir hyperlink ekler."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"📄 {label}: ")
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    rel_id = doc.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rPr.append(style)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.text = url
    run_el.append(t)
    hyperlink.append(run_el)
    p._p.append(hyperlink)


def add_summary_section(doc: Document, summary: dict) -> None:
    """3 maddelik Claude özetini rapora yazar."""
    items = [
        ("1. Eşyanın Ticari Tanımı",  summary.get("esya_tanimi", "-")),
        ("2. GTİP Kararı",            summary.get("gtip_karar", "-")),
        ("3. Teknik Gerekçe",         summary.get("teknik_gerekce", "-")),
    ]
    for title, content in items:
        tp = doc.add_paragraph()
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.size = Pt(11)
        tr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        cp = doc.add_paragraph()
        cp.add_run(content).font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph("─" * 80)
    doc.add_paragraph()


# ── Birleşik rapor bileşenleri ────────────────────────────────────────────────

def add_section_divider(doc: Document, title: str, bg_color: str = "1F4E79") -> None:
    """Ülke bölümü için renkli başlık çubuğu ekler."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def add_country_stats_card(
    doc: Document,
    countries: list[dict],   # [{"name": str, "count": int, "color": str}, ...]
) -> None:
    """
    Raporun başına renkli ülke kartlarını ekler.
    countries listesindeki her eleman bir renk kartı olarak gösterilir.
    """
    if not countries:
        return

    n = len(countries)
    widths = [int(9360 / n)] * n
    table = doc.add_table(rows=2, cols=n)
    table.style = "Table Grid"

    for i, c in enumerate(countries):
        # Üst satır: karar sayısı
        cell_top = table.rows[0].cells[i]
        set_cell_bg(cell_top, c.get("color", "1F4E79"))
        p = cell_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(c["count"]))
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Alt satır: ülke adı
        cell_bot = table.rows[1].cells[i]
        set_cell_bg(cell_bot, c.get("label_color", "2E75B6"))
        pb = cell_bot.paragraphs[0]
        pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = pb.add_run(c["name"])
        rb.bold = True
        rb.font.size = Pt(11)
        rb.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()
