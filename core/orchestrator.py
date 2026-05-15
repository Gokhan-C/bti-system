"""
Orchestrator — tüm connector'ları sırayla çalıştırır, birleşik rapor üretir.
"""

import importlib
import json
import sys
from datetime import datetime, timedelta
from pathlib import Path

import yaml

from core.base_connector import BaseConnector, ConnectorResult
from core.logger import get_logger
from core.report_builder import (
    make_doc, add_info_table, add_ruling_link,
    add_summary_section, add_no_results_notice,
    add_section_divider, add_country_stats_card, set_cell_bg,
)

from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Ülke renk paleti (birleşik rapor kartları)
COUNTRY_COLORS = {
    "eu_ebti":  {"color": "003399", "label_color": "2E75B6"},   # AB mavisi
    "us_cbp":   {"color": "B22234", "label_color": "C0392B"},   # ABD kırmızısı
    "ca_cbsa":  {"color": "FF0000", "label_color": "C0392B"},   # Kanada kırmızısı
}

COUNTRY_SECTION_COLORS = {
    "eu_ebti":  "1F3864",
    "us_cbp":   "7B241C",
    "ca_cbsa":  "922B21",
}


def _load_config(config_path: Path) -> dict:
    with open(config_path, encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    def expand(val):
        if isinstance(val, str):
            return str(Path(val).expanduser()) if val.startswith("~") else val
        if isinstance(val, dict):
            return {k: expand(v) for k, v in val.items()}
        return val

    return expand(cfg)


def _load_connectors(config, state_dir, output_base, logger) -> list[BaseConnector]:
    connectors = []
    for connector_id, connector_cfg in config.get("connectors", {}).items():
        if not connector_cfg.get("enabled", False):
            continue
        try:
            module     = importlib.import_module(f"connectors.{connector_id}")
            class_name = "".join(p.capitalize() for p in connector_id.split("_")) + "Connector"
            cls        = getattr(module, class_name)
            output_dir = output_base / connector_id.upper()
            if "claude_batch_size" not in connector_cfg and "claude_batch_size" in config:
                connector_cfg = {**connector_cfg, "claude_batch_size": config["claude_batch_size"]}
            connectors.append(cls(config=connector_cfg, state_dir=state_dir, output_dir=output_dir))
            logger.info(f"  Connector yüklendi: {connectors[-1].display_name}")
        except Exception as e:
            logger.error(f"  Connector yüklenemedi ({connector_id}): {e}")
    return connectors


# ── Cached report data yükleyici ─────────────────────────────────────────────

def _load_cached_report_data(connector, target_date: datetime, output_base: Path) -> dict:
    """
    result.data boş olduğunda (records_new=0 → build_report çağrılmadı),
    connector'ın output dizininde kayıtlı _report_data.json'ı bulup döndürür.
    """
    date_str = target_date.strftime("%Y-%m-%d")
    base_dir = output_base / connector.connector_id.upper()

    # 1. Hedef tarih dizini
    candidate = base_dir / date_str / "_report_data.json"
    if candidate.exists():
        try:
            return json.loads(candidate.read_text(encoding="utf-8"))
        except Exception:
            pass

    # 2. En son tarih dizini (fallback)
    try:
        date_dirs = sorted(
            [d for d in base_dir.iterdir() if d.is_dir() and (d / "_report_data.json").exists()],
            reverse=True,
        )
        if date_dirs:
            return json.loads((date_dirs[0] / "_report_data.json").read_text(encoding="utf-8"))
    except Exception:
        pass

    return {}


# ── Birleşik rapor ────────────────────────────────────────────────────────────

def _build_unified_report(
    results: list[ConnectorResult],
    connectors: list[BaseConnector],
    config: dict,
    target_date: datetime,
    output_base: Path,
    logger,
) -> Path | None:
    unified_cfg = config.get("unified_report", {})
    if not unified_cfg.get("enabled", False):
        return None

    date_str    = target_date.strftime("%Y-%m-%d")
    output_name = unified_cfg.get("output_name", "BTI_Unified_{date}.docx").replace("{date}", date_str)
    docx_path   = output_base / output_name

    connector_map = {c.connector_id: c for c in connectors}

    # ── Doküman başlangıcı ──────────────────────────────────────────────────
    doc = make_doc("", "")  # boş başlık, kapak bloğu elle oluşturulacak
    # make_doc'un eklediği boş paragrafları temizle
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.5))

    # ── Kapak: başlık ───────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("BAĞLAYICI TARİFE BİLGİLERİ")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x1F, 0x3E, 0x6E)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"Günlük Rapor  |  {date_str}")
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── Ülke kartları ────────────────────────────────────────────────────────
    result_map = {r.connector_id: r for r in results}
    country_cards = []
    for cid, conn in connector_map.items():
        res = result_map.get(cid)
        count = res.records_new if res else 0
        colors = COUNTRY_COLORS.get(cid, {"color": "444444", "label_color": "666666"})
        country_cards.append({
            "name":        conn.display_name,
            "count":       count,
            "color":       colors["color"],
            "label_color": colors["label_color"],
        })

    add_country_stats_card(doc, country_cards)

    # ── Toplam istatistik tablosu ─────────────────────────────────────────
    total_new = sum(r.records_new for r in results)
    stat_rows = [("Toplam Yeni Karar", str(total_new))]
    for res in results:
        conn = connector_map.get(res.connector_id)
        name = conn.display_name if conn else res.connector_id
        status = "✓" if res.success else "✗ HATA"
        stat_rows.append((name, f"{res.records_new} karar  {status}"))
    add_info_table(doc, stat_rows)
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Her ülkenin içeriği ────────────────────────────────────────────────
    for res in results:
        conn = connector_map.get(res.connector_id)
        if not conn or not res.success:
            continue

        section_color = COUNTRY_SECTION_COLORS.get(res.connector_id, "1F4E79")
        add_section_divider(doc, f"  {conn.display_name.upper()}  ", bg_color=section_color)

        data = res.data
        # result.data boşsa (records_new=0) → kalıcı JSON'dan yükle
        if not data:
            data = _load_cached_report_data(conn, target_date, output_base)
            if data:
                logger.info(f"  {conn.display_name}: cached _report_data.json yüklendi")

        # ── AB EBTI bölümü ──────────────────────────────────────────────
        if res.connector_id == "eu_ebti" and data:
            # Ülke dağılım özet tablosu
            country_stats = data.get("country_stats", [])
            if country_stats:
                p = doc.add_paragraph()
                r2 = p.add_run("Ülkelere Göre Dağılım")
                r2.bold = True
                r2.font.size = Pt(11)
                r2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                add_info_table(doc, [
                    (s["name"], f"{s['count']} BTI")
                    for s in country_stats
                ])
                doc.add_paragraph()

            # Kayıtlar
            for rec in data.get("records", []):
                ref  = rec.get("ref", "")
                hs   = rec.get("hs", "")
                date = rec.get("date_issue", "")
                desc = rec.get("desc_tr", "")
                just = rec.get("just_tr", "")
                url  = (
                    f"https://ec.europa.eu/taxation_customs/dds2/ebti/ebti_consultation.jsp"
                    f"?Lang=en&reference={ref.replace('/', '%2F')}&Expand=true&offset=1&allRecords=0"
                )
                add_ruling_link(doc, url, label=f"BTI Kararı {ref}")

                hp = doc.add_paragraph()
                hr2 = hp.add_run(f"GTİP {hs}  |  {ref}  |  {date}")
                hr2.bold = True
                hr2.font.size = Pt(12)
                hr2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

                add_summary_section(doc, {
                    "esya_tanimi":   desc,
                    "gtip_karar":    hs,
                    "teknik_gerekce": just,
                })

        # ── ABD CBP bölümü ──────────────────────────────────────────────
        elif res.connector_id == "us_cbp" and data:
            stats = data.get("stats", {})
            if stats:
                add_info_table(doc, [
                    ("Toplam Çekilen",          str(stats.get("total", 0))),
                    ("Tarife Sınıflandırması",  str(stats.get("classification", 0))),
                    ("Menşei Kararı",           str(stats.get("origin", 0))),
                    ("Diğer",                   str(stats.get("other", 0))),
                    ("Rapora Giren",            str(stats.get("in_report", 0))),
                ])
                doc.add_paragraph()

            for rec in data.get("records", []):
                add_ruling_link(doc, rec["source_url"])

                hp = doc.add_paragraph()
                hr2 = hp.add_run(f"ABD CBP: {rec['number']}  |  {rec['date_fmt']}")
                hr2.bold = True
                hr2.font.size = Pt(12)
                hr2.font.color.rgb = RGBColor(0x7B, 0x24, 0x1C)

                add_info_table(doc, [
                    ("Karar No",  rec["number"]),
                    ("GTİP",      rec["tariffs"]),
                    ("Tarih",     rec["date_fmt"]),
                ])
                doc.add_paragraph()
                add_summary_section(doc, rec.get("summary", {}))

        # ── Kanada CBSA bölümü ──────────────────────────────────────────
        elif res.connector_id == "ca_cbsa" and data:
            for rec in data.get("records", []):
                add_ruling_link(doc, rec["source_url"])

                hp = doc.add_paragraph()
                hr2 = hp.add_run(f"Kanada CBSA: {rec['ruling_id']}  |  {rec['date_fmt']}")
                hr2.bold = True
                hr2.font.size = Pt(12)
                hr2.font.color.rgb = RGBColor(0x92, 0x2B, 0x21)

                add_info_table(doc, [
                    ("Karar No",     rec["ruling_id"]),
                    ("GTİP",         rec["hts"]),
                    ("Tarih",        rec["date_fmt"]),
                    ("Başvurucu",    rec["applicant"]),
                    ("Menşe Ülke",   rec["origin"]),
                ])
                doc.add_paragraph()
                add_summary_section(doc, rec.get("summary", {}))

        if res.records_new == 0:
            add_no_results_notice(doc)
            doc.add_paragraph()

    doc.save(str(docx_path))
    logger.info(f"Birleşik rapor kaydedildi: {docx_path}")
    return docx_path


# ── Ana Orchestrator ──────────────────────────────────────────────────────────

class Orchestrator:

    def __init__(self, config_path: Path):
        self.config_path = Path(config_path)
        self.config      = _load_config(self.config_path)

        log_file    = self.config.get("log_file")
        self.logger = get_logger("bti_system", log_file=Path(log_file) if log_file else None)

        self.output_base = Path(self.config["output_base"])
        self.state_dir   = Path(self.config["state_dir"])
        self.output_base.mkdir(parents=True, exist_ok=True)
        self.state_dir.mkdir(parents=True, exist_ok=True)

        bti_root = str(self.config_path.parent)
        if bti_root not in sys.path:
            sys.path.insert(0, bti_root)

        self.connectors = _load_connectors(
            self.config, self.state_dir, self.output_base, self.logger
        )

    def run(self, target_date: datetime | None = None) -> list[ConnectorResult]:
        if target_date is None:
            target_date = datetime.now() - timedelta(days=1)

        self.logger.info("=" * 60)
        self.logger.info(f"BTI Sistemi | {target_date.strftime('%Y-%m-%d')}")
        self.logger.info("=" * 60)

        results: list[ConnectorResult] = []

        for connector in self.connectors:
            self.logger.info(f"── {connector.display_name} ──────────────")
            try:
                result = connector.run(target_date)
                results.append(result)
                if result.success:
                    self.logger.info(
                        f"  {connector.display_name}: {result.records_fetched} çekildi, "
                        f"{result.records_new} yeni → {[p.name for p in result.output_paths]}"
                    )
                else:
                    self.logger.error(f"  {connector.display_name} HATA:\n{''.join(result.errors)}")
            except Exception as e:
                import traceback
                self.logger.critical(f"  {connector.display_name} beklenmedik hata: {traceback.format_exc()}")
                results.append(ConnectorResult(
                    connector_id=connector.connector_id,
                    run_date=target_date,
                    records_fetched=0, records_new=0,
                    errors=[str(e)], success=False,
                ))

        self._print_summary(results)

        unified_path = _build_unified_report(
            results, self.connectors, self.config, target_date, self.output_base, self.logger
        )
        if unified_path:
            self.logger.info(f"Birleşik rapor: {unified_path}")

        self.logger.info("=" * 60)
        return results

    def _print_summary(self, results: list[ConnectorResult]) -> None:
        self.logger.info("")
        self.logger.info("┌─── ÖZET ───────────────────────────────────────┐")
        for r in results:
            status = "✓ BAŞARILI" if r.success else "✗ HATA"
            self.logger.info(
                f"│  {r.connector_id:<12} {status:<12} "
                f"çekilen:{r.records_fetched:>4}  yeni:{r.records_new:>4}  │"
            )
        self.logger.info("└────────────────────────────────────────────────┘")
        self.logger.info("")
